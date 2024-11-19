'''
文件说明：此模块负责处理文档上传和文本分割

支持PDF、DOCX和TXT格式文件的处理，使用LangChain的文本分割器
将文档分割成适合向量化的文本块
'''

from typing import List, Dict, Any
import os
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from utils.schema import Document

class DocumentProcessor:
    '''
    文档处理器类
    
    负责将上传的文档（PDF、DOCX、TXT）转换为文本块，
    并进行分割以便后续处理。支持多种文档格式的处理和
    自动文本分块。
    '''
    
    def __init__(self):
        '''
        初始化文档处理器
        
        配置文本分割器参数，包括块大小和重叠长度
        '''
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            length_function=len
        )

    def process_file(self, file, filename: str) -> List[Document]:
        '''
        处理上传的文件并返回文档块列表
        
        参数:
            file: 上传的文件对象
            filename: str - 文件名
            
        返回:
            List[Document] - 处理后的文档块列表
            
        异常:
            ValueError: 当文件格式不支持时抛出
            Exception: 处理文件时的其他错误
        '''
        try:
            # 获取文件扩展名
            _, ext = os.path.splitext(filename)
            ext = ext.lower()

            # 根据文件类型提取文本
            if ext == '.pdf':
                text = self._process_pdf(file)
            elif ext == '.docx':
                text = self._process_docx(file)
            elif ext == '.txt':
                text = file.read().decode('utf-8')#若为txt文件，则直接通过read()方法读取文件内容
            else:
                raise ValueError(f"不支持的文件类型：{ext}")

            # 分割文本为块
            chunks = self.text_splitter.split_text(text)

            # 创建Document对象
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    content=chunk,
                    metadata={
                        "filename": filename,
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)

            return documents

        except Exception as e:
            raise Exception(f"处理文件 {filename} 时出错：{str(e)}")

    def _process_pdf(self, file) -> str:
        '''
        处理PDF文件并提取文本
        
        参数:
            file: PDF文件对象
            
        返回:
            str - 提取的文本内容
            
        异常:
            Exception: PDF处理错误
        '''
        try:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"处理PDF时出错：{str(e)}")

    def _process_docx(self, file) -> str:
        '''
        处理DOCX文件并提取文本
        
        参数:
            file: DOCX文件对象
            
        返回:
            str - 提取的文本内容
            
        异常:
            Exception: DOCX处理错误
        '''
        try:
            doc = DocxDocument(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"处理DOCX时出错：{str(e)}")

    def process_document(self, text):
        # Assuming BAAI-BGE model is used for embedding
        # This function should preprocess the text for the model
        # Example: tokenization, removing special characters, etc.
        processed_text = text.lower().replace('[^\w\s]', '')
        #processed_text = re.sub(r'[^\w\s]', '', text.lower())
        return processed_text